"""
Document Q&A Assistant
=====================

A Python GUI application that allows users to load documents (TXT, DOCX, DOC, PDF)
and ask questions about their content using the Google Gemini AI API.

Features:
- Multiple document format support with fallback extraction methods
- Threaded operations to prevent GUI freezing
- Session management with save/clear functionality
- Dependency checking and user-friendly error handling
- Web search integration for comprehensive answers
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import requests
import os
import json
import datetime
import threading
from pathlib import Path

# =============================================================================
# DOCUMENT PROCESSING LIBRARIES WITH AVAILABILITY CHECKING
# =============================================================================
# Try to import various document processing libraries and set availability flags
# This approach allows graceful degradation if some libraries are not installed

LIBRARIES = {}  # Dictionary to track which libraries are available

try:
    from docx import Document
    LIBRARIES['docx'] = True
except ImportError:
    LIBRARIES['docx'] = False

try:
    import mammoth  # For Word document conversion
    LIBRARIES['mammoth'] = True
except ImportError:
    LIBRARIES['mammoth'] = False

try:
    import PyPDF2  # PDF processing library
    LIBRARIES['pypdf2'] = True
except ImportError:
    LIBRARIES['pypdf2'] = False

try:
    import pdfplumber  # Alternative PDF processing
    LIBRARIES['pdfplumber'] = True
except ImportError:
    LIBRARIES['pdfplumber'] = False

try:
    import fitz  # PyMuPDF - Another PDF library
    LIBRARIES['pymupdf'] = True
except ImportError:
    LIBRARIES['pymupdf'] = False


class DocumentQAApp:
    """
    Main application class for the Document Q&A Assistant.
    
    This class creates and manages the GUI interface, handles document processing,
    communicates with the Gemini API, and manages the Q&A session.
    
    Attributes:
        root (tk.Tk): The main Tkinter window
        document_path (tk.StringVar): Path to the currently loaded document
        document_content (str): Extracted text content from the document
        qa_pairs (list): History of question-answer pairs in the current session
        api_key (tk.StringVar): Gemini API key for authentication
    """
    
    def __init__(self, root):
        """
        Initialize the Document Q&A application.
        
        Args:
            root (tk.Tk): The main Tkinter window object
        """
        self.root = root
        self.root.title("Document Q&A Assistant")
        self.root.geometry("1000x700")
        self.root.minsize(800, 600)  # Set minimum window size
        
        # Application state variables
        self.document_path = tk.StringVar()  # Stores the path to the current document
        self.document_content = ""  # Stores the extracted text content
        self.qa_pairs = []  # Stores the history of Q&A pairs
        # Try to get API key from environment variable, default to empty string
        self.api_key = tk.StringVar(value=os.getenv("GEMINI_API_KEY", ""))
        
        # Initialize the GUI and check dependencies
        self.create_widgets()
        self.check_dependencies()
        
    def create_widgets(self):
        """
        Create and arrange all GUI widgets in the main window.
        
        This method sets up the main layout with sections for:
        - API key input
        - Document selection
        - Question input
        - Response display
        - Control buttons
        - Status bar
        """
        # Main container frame with padding
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights for responsive resizing
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)  # Second column expands
        main_frame.rowconfigure(4, weight=1)     # Response section expands
        
        # Create each section of the interface
        self._create_api_section(main_frame)        # API key input
        self._create_document_section(main_frame)   # Document selection
        self._create_question_section(main_frame)   # Question input
        self._create_response_section(main_frame)   # Response display
        self._create_control_buttons(main_frame)    # Control buttons
        self._create_status_bar(main_frame)         # Status bar
        
    def _create_api_section(self, parent):
        """
        Create the API key input section.
        
        Args:
            parent: The parent widget to attach this section to
        """
        ttk.Label(parent, text="Gemini API Key:").grid(
            row=0, column=0, sticky=tk.W, pady=(0, 5))
        
        # Frame to contain API entry and toggle button
        api_frame = ttk.Frame(parent)
        api_frame.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=(0, 5))
        api_frame.columnconfigure(0, weight=1)  # Entry field expands
        
        # API key entry field (masked by default)
        self.api_entry = ttk.Entry(api_frame, textvariable=self.api_key, 
                                  show="*", width=50)
        self.api_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        # Button to toggle visibility of the API key
        ttk.Button(api_frame, text="Show/Hide", 
                  command=self.toggle_api_visibility).grid(row=0, column=1)
    
    def _create_document_section(self, parent):
        """
        Create the document selection section.
        
        Args:
            parent: The parent widget to attach this section to
        """
        ttk.Label(parent, text="Document:").grid(
            row=1, column=0, sticky=tk.W, pady=(10, 5))
        
        # Frame for document path entry and buttons
        doc_frame = ttk.Frame(parent)
        doc_frame.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=(10, 5))
        doc_frame.columnconfigure(0, weight=1)  # Path entry expands
        
        # Read-only entry to display selected document path
        self.doc_entry = ttk.Entry(doc_frame, textvariable=self.document_path, 
                                  state="readonly")
        self.doc_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        
        # Buttons for browsing and loading documents
        ttk.Button(doc_frame, text="Browse", 
                  command=self.browse_document).grid(row=0, column=1)
        ttk.Button(doc_frame, text="Load", 
                  command=self.load_document).grid(row=0, column=2, padx=(5, 0))
        
        # Label to display document loading status and info
        self.info_label = ttk.Label(parent, text="No document loaded", 
                                   foreground="gray")
        self.info_label.grid(row=2, column=0, columnspan=2, sticky=tk.W, 
                            pady=(0, 10))
    
    def _create_question_section(self, parent):
        """
        Create the question input section.
        
        Args:
            parent: The parent widget to attach this section to
        """
        ttk.Label(parent, text="Ask a question:").grid(
            row=3, column=0, sticky=(tk.W, tk.N), pady=(0, 5))
        
        # Frame for question entry and ask button
        question_frame = ttk.Frame(parent)
        question_frame.grid(row=3, column=1, sticky=(tk.W, tk.E), pady=(0, 5))
        question_frame.columnconfigure(0, weight=1)  # Question entry expands
        
        # Entry field for typing questions
        self.question_entry = ttk.Entry(question_frame, font=("TkDefaultFont", 10))
        self.question_entry.grid(row=0, column=0, sticky=(tk.W, tk.E), padx=(0, 5))
        # Bind Enter key to trigger question submission
        self.question_entry.bind("<Return>", lambda e: self.ask_question())
        
        # Ask button (initially disabled until document is loaded)
        self.ask_button = ttk.Button(question_frame, text="Ask", 
                                    command=self.ask_question, state="disabled")
        self.ask_button.grid(row=0, column=1)
    
    def _create_response_section(self, parent):
        """
        Create the response display section with a scrollable text area.
        
        Args:
            parent: The parent widget to attach this section to
        """
        # Labeled frame for the Q&A session
        response_frame = ttk.LabelFrame(parent, text="Q&A Session", padding="5")
        response_frame.grid(row=4, column=0, columnspan=2, 
                           sticky=(tk.W, tk.E, tk.N, tk.S), pady=(10, 0))
        response_frame.columnconfigure(0, weight=1)
        response_frame.rowconfigure(0, weight=1)
        
        # Scrollable text area for displaying Q&A
        self.response_text = scrolledtext.ScrolledText(
            response_frame, 
            wrap=tk.WORD,  # Wrap text at word boundaries
            font=("Consolas", 10),  # Monospaced font for better readability
            state="disabled"  # Initially disabled to prevent direct editing
        )
        self.response_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure text formatting tags for different types of content
        tags = {
            "question": {"foreground": "blue", "font": ("TkDefaultFont", 10, "bold")},
            "answer": {"foreground": "black", "font": ("TkDefaultFont", 10)},
            "info": {"foreground": "gray", "font": ("TkDefaultFont", 9)},
            "error": {"foreground": "red", "font": ("TkDefaultFont", 10)}
        }
        for tag, config in tags.items():
            self.response_text.tag_config(tag, **config)
    
    def _create_control_buttons(self, parent):
        """
        Create control buttons for session management.
        
        Args:
            parent: The parent widget to attach these buttons to
        """
        button_frame = ttk.Frame(parent)
        button_frame.grid(row=5, column=0, columnspan=2, pady=(10, 0))
        
        # Define buttons with their text and command functions
        buttons = [
            ("Clear Session", self.clear_session),
            ("Save Session", self.save_session),
            ("Check Dependencies", self.show_dependencies)
        ]
        
        # Create and pack each button
        for i, (text, command) in enumerate(buttons):
            ttk.Button(button_frame, text=text, command=command).pack(
                side=tk.LEFT, padx=(0, 5))
    
    def _create_status_bar(self, parent):
        """
        Create a status bar at the bottom of the window.
        
        Args:
            parent: The parent widget to attach the status bar to
        """
        self.status_var = tk.StringVar(value="Ready")  # Variable for status text
        status_bar = ttk.Label(parent, textvariable=self.status_var, 
                              relief=tk.SUNKEN, anchor=tk.W)
        status_bar.grid(row=6, column=0, columnspan=2, sticky=(tk.W, tk.E), 
                       pady=(5, 0))
    
    def _update_ui_thread_safe(self, callback):
        """
        Helper method to safely update UI from background threads.
        
        Tkinter is not thread-safe, so UI updates from background threads
        must be scheduled to run in the main thread using the after() method.
        
        Args:
            callback: A function to be executed in the main thread
        """
        self.root.after(0, callback)
    
    def _show_error(self, message):
        """
        Helper method to display error messages.
        
        Args:
            message (str): The error message to display
        """
        messagebox.showerror("Error", message)
    
    def toggle_api_visibility(self):
        """Toggle the visibility of the API key in the entry field."""
        # Switch between showing asterisks or plain text
        show_state = "" if self.api_entry.cget("show") == "*" else "*"
        self.api_entry.config(show=show_state)
    
    def check_dependencies(self):
        """Check available document processing libraries and update status."""
        # Get list of available libraries
        available_libs = [name for name, available in LIBRARIES.items() if available]
        if available_libs:
            self.add_to_response(f"Available libraries: {', '.join(available_libs)}\n", "info")
        else:
            self.add_to_response("Warning: No document processing libraries found!\n", "error")
    
    def show_dependencies(self):
        """
        Display a dialog with information about available and missing dependencies.
        
        This helps users understand what functionality is available and
        what they need to install for full support.
        """
        # Library information: (library name, available message, missing message)
        lib_info = {
            'docx': ('python-docx', 'Available (.docx files)', 'Not installed (.docx files)'),
            'mammoth': ('mammoth', 'Available (.docx/.doc files)', 'Not installed (.docx/.doc files)'),
            'pdfplumber': ('pdfplumber', 'Available (.pdf files)', 'Not installed (.pdf files)'),
            'pymupdf': ('PyMuPDF', 'Available (.pdf files)', 'Not installed (.pdf files)'),
            'pypdf2': ('PyPDF2', 'Available (.pdf files)', 'Not installed (.pdf files)')
        }
        
        # Build dependency information string
        deps_info = "Document Processing Libraries:\n\n"
        deps_info += "✓ Text files (.txt): Built-in support\n"
        
        # Add status for each library
        for lib_key, (lib_name, available_msg, missing_msg) in lib_info.items():
            status = "✓" if LIBRARIES[lib_key] else "✗"
            message = available_msg if LIBRARIES[lib_key] else missing_msg
            deps_info += f"{status} {lib_name}: {message}\n"
        
        # Add installation commands
        deps_info += "\n\nInstallation commands:\n"
        deps_info += "pip install python-docx mammoth\n"
        deps_info += "pip install pdfplumber PyMuPDF PyPDF2"
        
        # Show information dialog
        messagebox.showinfo("Dependencies", deps_info)
    
    def browse_document(self):
        """Open a file dialog to select a document file."""
        # Define supported file types
        filetypes = [
            ("All supported", "*.txt;*.docx;*.doc;*.pdf"),
            ("Text files", "*.txt"),
            ("Word documents", "*.docx;*.doc"),
            ("PDF files", "*.pdf"),
            ("All files", "*.*")
        ]
        
        # Show file dialog and update document path if a file is selected
        filename = filedialog.askopenfilename(title="Select Document", filetypes=filetypes)
        if filename:
            self.document_path.set(filename)
    
    def load_document(self):
        """Load the selected document and extract its text content."""
        doc_path = self.document_path.get()
        if not doc_path:
            self._show_error("Please select a document first.")
            return
        
        if not os.path.exists(doc_path):
            self._show_error("Selected file does not exist.")
            return
        
        # Start document loading in a background thread to prevent UI freezing
        threading.Thread(target=self._load_document_thread, daemon=True).start()
    
    def _load_document_thread(self):
        """
        Background thread for loading document content.
        
        This method runs in a separate thread to prevent the UI from freezing
        during document processing, which can be time-consuming for large files.
        """
        try:
            # Update UI to show loading status
            self._update_ui_thread_safe(lambda: self.status_var.set("Loading document..."))
            self._update_ui_thread_safe(lambda: self.ask_button.config(state="disabled"))
            
            # Extract text content from the document
            content = self.read_document_content(self.document_path.get())
            
            if content and content.strip():
                # Successfully extracted content
                self.document_content = content
                filename = os.path.basename(self.document_path.get())
                char_count = len(content)
                
                # Update UI with document information
                info_text = f"Loaded: {filename} ({char_count:,} characters)"
                self._update_ui_thread_safe(lambda: self.info_label.config(
                    text=info_text, foreground="green"))
                self._update_ui_thread_safe(lambda: self.ask_button.config(state="normal"))
                self._update_ui_thread_safe(lambda: self.question_entry.focus())
                
                # Add loading information to response area
                self._update_ui_thread_safe(lambda: self.add_to_response(
                    f"Document loaded: {filename}\n", "info"))
                # Show a preview of the document content
                preview = content[:200] + ('...' if len(content) > 200 else '')
                self._update_ui_thread_safe(lambda: self.add_to_response(
                    f"Content preview:\n{preview}\n\n", "info"))
                self._update_ui_thread_safe(lambda: self.status_var.set("Document loaded successfully"))
            else:
                # Failed to extract content
                self._update_ui_thread_safe(lambda: self._show_error(
                    "Failed to extract text from document or document is empty."))
                self._update_ui_thread_safe(lambda: self.status_var.set("Failed to load document"))
                
        except Exception as e:
            # Handle any errors during document loading
            self._update_ui_thread_safe(lambda: self._show_error(
                f"Error loading document: {str(e)}"))
            self._update_ui_thread_safe(lambda: self.status_var.set("Error loading document"))
    
    def ask_question(self):
        """Process a question about the loaded document."""
        # Validate inputs
        if not self.api_key.get().strip():
            self._show_error("Please enter your Gemini API key.")
            return
        
        if not self.document_content:
            self._show_error("Please load a document first.")
            return
        
        question = self.question_entry.get().strip()
        if not question:
            messagebox.showwarning("Warning", "Please enter a question.")
            return
        
        # Clear the question entry and display the question
        self.question_entry.delete(0, tk.END)
        self.add_to_response(f"Q: {question}\n", "question")
        
        # Process the question in a background thread
        threading.Thread(target=self._ask_question_thread, args=(question,), daemon=True).start()
    
    def _ask_question_thread(self, question):
        """
        Background thread for processing questions with the Gemini API.
        
        Args:
            question (str): The question to ask about the document
        """
        try:
            # Update UI to show processing status
            self._update_ui_thread_safe(lambda: self.status_var.set("Processing question..."))
            self._update_ui_thread_safe(lambda: self.ask_button.config(state="disabled"))
            
            # Send question to Gemini API
            result = self.send_question_to_gemini(
                self.document_content, question, self.api_key.get())
            
            if result:
                generated_parts, search_queries = result
                
                # Display the answer
                self._update_ui_thread_safe(lambda: self.add_to_response("A: ", "question"))
                
                # Add each part of the generated response
                for part in generated_parts:
                    if "text" in part:
                        self._update_ui_thread_safe(lambda text=part["text"]: 
                            self.add_to_response(f"{text}\n", "answer"))
                
                # Show web search queries if any were considered
                if search_queries:
                    queries_text = ', '.join(search_queries)
                    self._update_ui_thread_safe(lambda: self.add_to_response(
                        f"Web searches considered: {queries_text}\n", "info"))
                
                # Add separator between Q&A pairs
                self._update_ui_thread_safe(lambda: self.add_to_response(
                    "\n" + "-"*50 + "\n\n", "info"))
                
                # Store the Q&A pair for session management
                response_texts = [part["text"] for part in generated_parts if "text" in part]
                self.qa_pairs.append({
                    'question': question,
                    'response_parts': response_texts,
                    'search_queries': search_queries if isinstance(search_queries, list) else []
                })
                
                self._update_ui_thread_safe(lambda: self.status_var.set("Question answered"))
            else:
                # Handle API failure
                self._update_ui_thread_safe(lambda: self.add_to_response(
                    "Error: Failed to get response from Gemini API.\n\n", "error"))
                self._update_ui_thread_safe(lambda: self.status_var.set("Error processing question"))
                
        except Exception as e:
            # Handle any errors during question processing
            self._update_ui_thread_safe(lambda: self.add_to_response(
                f"Error: {str(e)}\n\n", "error"))
            self._update_ui_thread_safe(lambda: self.status_var.set("Error processing question"))
        finally:
            # Re-enable the ask button regardless of success or failure
            self._update_ui_thread_safe(lambda: self.ask_button.config(state="normal"))
    
    def add_to_response(self, text, tag="answer"):
        """
        Add text to the response area with specified formatting.
        
        Args:
            text (str): The text to add
            tag (str): The formatting tag to apply (question, answer, info, error)
        """
        self.response_text.config(state="normal")  # Enable editing
        self.response_text.insert(tk.END, text, tag)  # Insert text with tag
        self.response_text.see(tk.END)  # Scroll to the bottom
        self.response_text.config(state="disabled")  # Disable editing
    
    def clear_session(self):
        """Clear the current Q&A session after confirmation."""
        if messagebox.askyesno("Clear Session", "Are you sure you want to clear the current session?"):
            self.response_text.config(state="normal")
            self.response_text.delete(1.0, tk.END)  # Clear all text
            self.response_text.config(state="disabled")
            self.qa_pairs.clear()  # Clear Q&A history
            self.status_var.set("Session cleared")
    
    def save_session(self):
        """Save the current Q&A session to a file."""
        if not self.qa_pairs:
            messagebox.showwarning("Warning", "No Q&A session to save.")
            return
        
        # Show file dialog to choose save location
        filename = filedialog.asksaveasfilename(
            title="Save Q&A Session",
            defaultextension=".md",  # Default to Markdown format
            filetypes=[("Markdown files", "*.md"), ("All files", "*.*")]
        )
        
        if filename:
            try:
                # Save the session
                self.save_qa_session(self.document_path.get(), self.document_content, 
                                   self.qa_pairs, filename)
                messagebox.showinfo("Success", f"Session saved to {filename}")
                self.status_var.set("Session saved")
            except Exception as e:
                self._show_error(f"Failed to save session: {str(e)}")
    
    def read_document_content(self, filepath):
        """
        Read text content from a document file based on its extension.
        
        Args:
            filepath (str): Path to the document file
            
        Returns:
            str: Extracted text content from the document
            
        Raises:
            Exception: If the file format is unsupported or extraction fails
        """
        # Get file extension to determine the appropriate reader
        file_ext = os.path.splitext(filepath)[1].lower()
        
        # Map file extensions to reader methods
        readers = {
            ".txt": self.read_text_file,
            ".docx": self._read_docx_file,
            ".doc": self._read_doc_file,
            ".pdf": self.extract_text_from_pdf
        }
        
        # Get the appropriate reader function
        reader = readers.get(file_ext)
        if not reader:
            raise Exception(f"Unsupported file format: {file_ext}")
        
        # Extract content using the appropriate reader
        content = reader(filepath)
        if not content or not content.strip():
            raise Exception(f"Failed to extract text from {file_ext} file or file is empty")
        
        return content
    
    def _read_docx_file(self, filepath):
        """
        Read .docx file with multiple fallback extraction methods.
        
        Args:
            filepath (str): Path to the .docx file
            
        Returns:
            str: Extracted text content or None if extraction fails
        """
        # Try python-docx first
        if LIBRARIES['docx']:
            content = self.extract_text_from_docx(filepath)
            if content and content.strip():
                return content
        
        # Fallback to mammoth if python-docx fails or is not available
        if LIBRARIES['mammoth']:
            content = self.extract_text_with_mammoth(filepath)
            if content and content.strip():
                return content
        
        return None
    
    def _read_doc_file(self, filepath):
        """
        Read .doc file using mammoth library.
        
        Args:
            filepath (str): Path to the .doc file
            
        Returns:
            str: Extracted text content or None if extraction fails
        """
        if LIBRARIES['mammoth']:
            return self.extract_text_with_mammoth(filepath)
        return None
    
    def read_text_file(self, filepath):
        """
        Read a plain text file with multiple encoding attempts.
        
        Args:
            filepath (str): Path to the text file
            
        Returns:
            str: Content of the text file
            
        Raises:
            Exception: If no supported encoding works
        """
        # Try multiple encodings to handle various text file formats
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(filepath, "r", encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue  # Try next encoding
        
        raise Exception("Failed to read text file with any supported encoding")
    
    def extract_text_from_docx(self, filepath):
        """
        Extract text from a .docx file using python-docx library.
        
        Args:
            filepath (str): Path to the .docx file
            
        Returns:
            str: Extracted text content
        """
        doc = Document(filepath)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
    
    def extract_text_with_mammoth(self, filepath):
        """
        Extract text from a Word file using mammoth library.
        
        Args:
            filepath (str): Path to the Word file
            
        Returns:
            str: Extracted text content
        """
        with open(filepath, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            return result.value
    
    def extract_text_from_pdf(self, filepath):
        """
        Extract text from PDF using available libraries with fallback.
        
        Args:
            filepath (str): Path to the PDF file
            
        Returns:
            str: Extracted text content
            
        Raises:
            Exception: If all extraction methods fail
        """
        # Try PDF extraction libraries in order of preference
        pdf_readers = [
            ('pdfplumber', self.extract_text_from_pdf_pdfplumber),
            ('pymupdf', self.extract_text_from_pdf_pymupdf),
            ('pypdf2', self.extract_text_from_pdf_pypdf2)
        ]
        
        for lib_name, reader_func in pdf_readers:
            if LIBRARIES[lib_name]:
                try:
                    return reader_func(filepath)
                except:
                    continue  # Try next library if this one fails
        
        raise Exception("Failed to extract PDF text with available libraries")
    
    def extract_text_from_pdf_pdfplumber(self, filepath):
        """
        Extract text from PDF using pdfplumber library.
        
        Args:
            filepath (str): Path to the PDF file
            
        Returns:
            str: Extracted text content with page markers
        """
        import pdfplumber
        text_content = []
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
        return "\n\n".join(text_content)
    
    def extract_text_from_pdf_pymupdf(self, filepath):
        """
        Extract text from PDF using PyMuPDF library.
        
        Args:
            filepath (str): Path to the PDF file
            
        Returns:
            str: Extracted text content with page markers
        """
        text_content = []
        pdf_document = fitz.open(filepath)
        try:
            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                page_text = page.get_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
        finally:
            pdf_document.close()  # Ensure file is closed properly
        return "\n\n".join(text_content)
    
    def extract_text_from_pdf_pypdf2(self, filepath):
        """
        Extract text from PDF using PyPDF2 library.
        
        Args:
            filepath (str): Path to the PDF file
            
        Returns:
            str: Extracted text content with page markers
        """
        text_content = []
        with open(filepath, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(pdf_reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
        return "\n\n".join(text_content)
    
    def send_question_to_gemini(self, document_content, user_question, api_key):
        """
        Send a question about the document to the Gemini API.
        
        Args:
            document_content (str): The text content of the document
            user_question (str): The question to ask about the document
            api_key (str): Gemini API key for authentication
            
        Returns:
            tuple: (generated_parts, search_queries) from the API response
            
        Raises:
            Exception: If the API request fails
        """
        # System instruction for the AI model
        instruction = """
You are a helpful AI assistant with a flourishing mindset.
Your primary goal is to answer questions based on the provided document content.
If the answer cannot be found in the document, state that the information is not present in the document.
"""

        # Construct the prompt with document content and question
        text = f"""
Here is the document content extracted from the file:

---
{document_content}
---

Based *only* on the document content provided above, please answer the following question:

Question: {user_question}
"""

        # API request headers
        headers = {
            "x-goog-api-key": api_key,
            "Content-Type": "application/json"
        }

        # API request payload
        request_payload = {
            "system_instruction": {"parts": [{"text": instruction}]},
            "contents": [{"parts": [{"text": text}]}],
            "generationConfig": {
                "temperature": 0.5,  # Controls randomness (0=deterministic, 1=creative)
                "topP": 0.95,       # Nucleus sampling parameter
                "topK": 64,         # Top-k sampling parameter
                "thinkingConfig": {
                    "thinkingBudget": -1,  # Unlimited thinking time
                    "includeThoughts": True  # Include reasoning process
                }
            },
            "tools": [{"google_search": {}}]  # Enable web search integration
        }

        try:
            # Send request to Gemini API
            response = requests.post(
                "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent",
                headers=headers,
                json=request_payload,
                timeout=240  # 4 minute timeout for large documents
            )
            response.raise_for_status()  # Raise exception for HTTP errors
            response_data = response.json()
            
            # Check for API errors
            if "error" in response_data:
                raise Exception(f"API Error: {response_data['error']}")
            
            # Extract response content and search queries
            generated_parts = response_data["candidates"][0]["content"]["parts"]
            search_queries = response_data["candidates"][0].get(
                "groundingMetadata", {}).get("webSearchQueries", [])
            
            return generated_parts, search_queries
            
        except Exception as e:
            raise Exception(f"Request failed: {str(e)}")
    
    def save_qa_session(self, document_path, document_content, qa_pairs, output_path):
        """
        Save the Q&A session to a Markdown file.
        
        Args:
            document_path (str): Path to the original document
            document_content (str): Content of the document
            qa_pairs (list): List of Q&A pairs from the session
            output_path (str): Path where to save the session file
        """
        with open(output_path, "w", encoding="utf-8") as f:
            # Write header information
            f.write("# Gemini API Document Q&A Session\n\n")
            f.write(f"**Date/Time:** {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("---\n\n")

            # Write document information
            f.write("## Document Information\n")
            f.write(f"**Selected Document:** `{os.path.basename(document_path)}`\n")
            f.write(f"**Document Type:** {os.path.splitext(document_path)[1].upper()} file\n")
            f.write(f"**Content Length:** {len(document_content)} characters\n")
            f.write(f"**Total Questions Asked:** {len(qa_pairs)}\n\n")
            f.write("---\n\n")

            # Write each Q&A pair
            for i, qa in enumerate(qa_pairs, 1):
                f.write(f"## Question {i}\n\n")
                f.write(f"**Question:** {qa['question']}\n\n")
                f.write("### Answer\n\n")
                
                # Write each part of the answer
                for part_text in qa['response_parts']:
                    f.write(part_text + "\n\n")
                
                # Write web search queries if any
                f.write("### Web Search Queries Considered\n")
                if qa['search_queries']:
                    for q in qa['search_queries']:
                        f.write(f"- {q}\n")
                    f.write("\n")
                else:
                    f.write("No web search queries were considered for this question.\n\n")
                
                f.write("---\n\n")


def main():
    """
    Main function to initialize and run the GUI application.
    """
    # Create the main window
    root = tk.Tk()
    app = DocumentQAApp(root)
    
    # Center the window on the screen
    root.update_idletasks()  # Ensure window dimensions are calculated
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")
    
    # Start the Tkinter event loop
    root.mainloop()


if __name__ == "__main__":
    main()